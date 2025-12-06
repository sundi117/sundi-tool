# -*- coding: utf-8 -*-
"""
@Description :  脚本： 生成数据字典（GUI版本）
@Author : sundi
@Created  : 2025/9/15 13:38
"""

import pymysql
import tkinter as tk
from tkinter import filedialog, messagebox
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import threading
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ==================== 数据库配置 ====================
# DB_HOST = '192.168.0.195'
# DB_USER = 'root'
# DB_PASSWORD = 'Ytzn@888#2022#@'
# DB_NAME = 'jiangsu-query'
# DB_CHARSET = 'utf8mb4'
# ==================================================

def connect_db(host, user, password, database=None, charset='utf8mb4'):
    """连接MySQL数据库"""
    return pymysql.connect(
        host=host,
        user=user,
        password=password,
        database=database,
        charset=charset
    )


def get_databases(host, user, password, charset='utf8mb4'):
    """获取MySQL服务器中所有数据库列表"""
    conn = connect_db(host, user, password, None, charset)
    cursor = conn.cursor()
    cursor.execute("SHOW DATABASES")
    databases = [db[0] for db in cursor.fetchall() if db[0] not in ['information_schema', 'performance_schema', 'mysql', 'sys']]
    cursor.close()
    conn.close()
    return databases


def get_tables(cursor, database_name):
    """获取数据库中所有表名"""
    cursor.execute("""
                   SELECT TABLE_NAME, TABLE_COMMENT
                   FROM information_schema.TABLES
                   WHERE TABLE_SCHEMA = %s
                     AND TABLE_TYPE = 'BASE TABLE'
                   """, (database_name,))
    return cursor.fetchall()


def get_columns(cursor, database_name, table_name):
    """获取指定表的所有字段信息"""
    cursor.execute("""
                   SELECT COLUMN_NAME,
                          COLUMN_TYPE,
                          IS_NULLABLE,
                          COLUMN_DEFAULT,
                          COLUMN_COMMENT,
                          EXTRA
                   FROM information_schema.COLUMNS
                   WHERE TABLE_SCHEMA = %s
                     AND TABLE_NAME = %s
                   ORDER BY ORDINAL_POSITION
                   """, (database_name, table_name))
    return cursor.fetchall()


def generate_markdown(cursor, database_name, tables, output_path):
    """生成Markdown格式数据字典"""
    file_path = os.path.join(output_path, f"数据字典-{database_name}.md")
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(f"# {database_name} 数据字典\n\n")
        for table in tables:
            table_name, table_comment = table
            columns = get_columns(cursor, database_name, table_name)

            f.write(f"## {table_name} \n")
            f.write(f"> **表注释**: {table_comment or '无'}\n\n")
            f.write("| 字段名 | 数据类型 | 是否可空 | 默认值 | 备注 | 额外信息 |\n")
            f.write("|--------|----------|----------|--------|------|----------|\n")

            for col in columns:
                col_name, col_type, nullable, default, comment, extra = col
                nullable = "是" if nullable == "YES" else "否"
                default = f"`{default}`" if default else "-"
                f.write(f"| `{col_name}` | `{col_type}` | {nullable} | {default} | {comment or '-'} | {extra or '-'} |\n")

            f.write("\n---\n\n")
    return file_path


def generate_doc(cursor, database_name, tables, output_path):
    """生成Word文档格式数据字典"""
    file_path = os.path.join(output_path, f"数据字典-{database_name}.docx")
    doc = Document()

    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(10.5)

    # 添加标题
    title_para = doc.add_paragraph(f'{database_name} 数据字典')
    title_para.style = 'Title'
    title_run = title_para.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.name = '微软雅黑'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 遍历每个表
    for table in tables:
        table_name, table_comment = table
        columns = get_columns(cursor, database_name, table_name)

        # 添加表名作为二级标题
        heading = doc.add_heading(table_name, level=1)

        # 添加表注释
        if table_comment:
            comment_para = doc.add_paragraph(f'表注释：{table_comment}')
            comment_para.italic = True
        else:
            comment_para = doc.add_paragraph('表注释：无')
            comment_para.italic = True

        # 创建表格
        table_doc = doc.add_table(rows=1, cols=6)
        table_doc.style = 'Light Grid Accent 1'

        # 设置表头
        header_cells = table_doc.rows[0].cells
        headers = ['字段名', '数据类型', '是否可空', '默认值', '备注', '额外信息']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # 设置表头字体加粗
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 添加数据行
        for col in columns:
            col_name, col_type, nullable, default, comment, extra = col
            nullable = "是" if nullable == "YES" else "否"
            default = str(default) if default else "-"
            comment = comment or "-"
            extra = extra or "-"

            row_cells = table_doc.add_row().cells
            row_cells[0].text = col_name
            row_cells[1].text = col_type
            row_cells[2].text = nullable
            row_cells[3].text = default
            row_cells[4].text = comment
            row_cells[5].text = extra

            # 设置单元格字体
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 添加分页符（除了最后一个表）
        if table != tables[-1]:
            doc.add_page_break()

    # 保存文档
    doc.save(file_path)
    return file_path


class DatabaseDictGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("数据库字典生成工具")
        self.root.geometry("780x920")
        self.root.resizable(False, False)

        # 输出路径
        self.output_path = os.getcwd()

        # 设置窗口居中
        self.center_window()

        # 创建界面
        self.create_widgets()

    def center_window(self):
        """窗口居中显示"""
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')

    def create_widgets(self):
        """创建界面组件"""
        # 主框架
        main_frame = ttk.Frame(self.root, padding=30)
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 标题区域
        title_frame = ttk.Frame(main_frame)
        title_frame.pack(fill=tk.X, pady=(0, 30))

        # 标题容器（包含标题和按钮，整体居中）
        title_container = ttk.Frame(title_frame)
        title_container.pack()

        title_label = ttk.Label(
            title_container,
            text="数据库字典生成工具",
            font=('微软雅黑', 20, 'bold'),
            bootstyle=PRIMARY
        )
        title_label.pack(side=tk.LEFT)

        # 信息按钮（紧靠标题）
        info_button = ttk.Button(
            title_container,
            text="ℹ️",
            command=self.show_about,
            bootstyle=OUTLINE,
            width=1
        )
        info_button.pack(side=tk.LEFT, padx=(8, 0))

        subtitle_label = ttk.Label(
            title_frame,
            text="快速生成MySQL数据库结构文档",
            font=('微软雅黑', 10),
            bootstyle=SECONDARY
        )
        subtitle_label.pack(pady=(8, 0))

        # 数据库配置框架
        config_frame = ttk.Labelframe(
            main_frame,
            text="📊 数据库配置",
            padding=20,
            bootstyle=INFO
        )
        config_frame.pack(fill=tk.X, pady=(0, 20))

        # 主机地址
        host_label = ttk.Label(config_frame, text="主机地址", font=('微软雅黑', 10))
        host_label.grid(row=0, column=0, sticky=tk.W, pady=12, padx=(0, 15))
        self.host_entry = ttk.Entry(config_frame, width=35, font=('微软雅黑', 10))
        self.host_entry.grid(row=0, column=1, pady=12, padx=5, sticky=tk.W + tk.E)
        self.host_entry.insert(0, "localhost")
        config_frame.grid_columnconfigure(1, weight=1)

        # 用户名
        user_label = ttk.Label(config_frame, text="用户名", font=('微软雅黑', 10))
        user_label.grid(row=1, column=0, sticky=tk.W, pady=12, padx=(0, 15))
        self.user_entry = ttk.Entry(config_frame, width=35, font=('微软雅黑', 10))
        self.user_entry.grid(row=1, column=1, pady=12, padx=5, sticky=tk.W + tk.E)
        self.user_entry.insert(0, "root")

        # 密码
        password_label = ttk.Label(config_frame, text="密码", font=('微软雅黑', 10))
        password_label.grid(row=2, column=0, sticky=tk.W, pady=12, padx=(0, 15))
        self.password_entry = ttk.Entry(config_frame, width=35, show="*", font=('微软雅黑', 10))
        self.password_entry.grid(row=2, column=1, pady=12, padx=5, sticky=tk.W + tk.E)

        # 数据库名（下拉菜单）
        db_label = ttk.Label(config_frame, text="数据库名", font=('微软雅黑', 10))
        db_label.grid(row=3, column=0, sticky=tk.W, pady=12, padx=(0, 15))
        database_frame = ttk.Frame(config_frame)
        database_frame.grid(row=3, column=1, sticky=tk.W + tk.E, pady=12, padx=5)
        database_frame.grid_columnconfigure(0, weight=1)

        self.database_combo = ttk.Combobox(
            database_frame,
            width=28,
            state="readonly",
            font=('微软雅黑', 10),
            bootstyle=INFO
        )
        self.database_combo.grid(row=0, column=0, sticky=tk.W + tk.E, padx=(0, 10))

        self.refresh_db_button = ttk.Button(
            database_frame,
            text="🔄 刷新列表",
            command=self.refresh_databases,
            bootstyle=OUTLINE,
            width=14
        )
        self.refresh_db_button.grid(row=0, column=1, sticky=tk.W)

        # 字符集
        # charset_label = ttk.Label(config_frame, text="字符集", font=('微软雅黑', 10))
        # charset_label.grid(row=4, column=0, sticky=tk.W, pady=12, padx=(0, 15))
        self.charset_entry = ttk.Entry(config_frame, width=35, font=('微软雅黑', 10))
        # self.charset_entry.grid(row=4, column=1, pady=12, padx=5, sticky=tk.W+tk.E)
        self.charset_entry.insert(0, "utf8mb4")

        # 输出路径框架
        path_frame = ttk.Labelframe(
            main_frame,
            text="📁 输出路径",
            padding=20,
            bootstyle=INFO
        )
        path_frame.pack(fill=tk.X, pady=(0, 20))

        # 格式选择
        format_label = ttk.Label(path_frame, text="输出格式", font=('微软雅黑', 10))
        format_label.pack(anchor=tk.W, pady=(0, 8))

        format_frame = ttk.Frame(path_frame)
        format_frame.pack(fill=tk.X, pady=(0, 15))

        self.format_combo = ttk.Combobox(
            format_frame,
            values=["Markdown格式", "Word文档格式"],
            state="readonly",
            font=('微软雅黑', 10),
            bootstyle=INFO,
            width=20
        )
        self.format_combo.pack(side=tk.LEFT)
        self.format_combo.current(0)  # 默认选择Markdown格式

        path_label = ttk.Label(path_frame, text="保存路径", font=('微软雅黑', 10))
        path_label.pack(anchor=tk.W, pady=(0, 10))

        path_input_frame = ttk.Frame(path_frame)
        path_input_frame.pack(fill=tk.X)
        path_input_frame.grid_columnconfigure(0, weight=1)

        self.path_entry = ttk.Entry(path_input_frame, font=('微软雅黑', 10))
        self.path_entry.grid(row=0, column=0, sticky=tk.W + tk.E, padx=(0, 10))
        self.path_entry.insert(0, self.output_path)

        browse_button = ttk.Button(
            path_input_frame,
            text="📂 浏览",
            command=self.select_path,
            bootstyle=OUTLINE,
            width=14
        )
        browse_button.grid(row=0, column=1, sticky=tk.W)

        # 状态显示区域
        status_frame = ttk.Frame(main_frame)
        status_frame.pack(fill=tk.X, pady=(0, 20))

        status_title = ttk.Label(
            status_frame,
            text="状态：",
            font=('微软雅黑', 10, 'bold')
        )
        status_title.pack(side=tk.LEFT, padx=(0, 10))

        self.status_label = ttk.Label(
            status_frame,
            text="✓ 就绪",
            font=('微软雅黑', 10),
            bootstyle=SUCCESS
        )
        self.status_label.pack(side=tk.LEFT)

        # 执行按钮区域
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X)

        self.generate_button = ttk.Button(
            button_frame,
            text="✨ 生成数据字典",
            command=self.start_generate,
            bootstyle=PRIMARY,
            width=30
        )
        self.generate_button.pack(pady=5)

    def show_about(self):
        """显示关于信息"""
        about_window = ttk.Toplevel(self.root)
        about_window.title("关于")
        about_window.geometry("500x400")
        about_window.resizable(False, False)

        # 居中显示
        about_window.update_idletasks()
        x = (about_window.winfo_screenwidth() // 2) - (500 // 2)
        y = (about_window.winfo_screenheight() // 2) - (400 // 2)
        about_window.geometry(f'500x400+{x}+{y}')

        # 主框架
        main_frame = ttk.Frame(about_window, padding=30)
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 标题
        title_label = ttk.Label(
            main_frame,
            text="数据库字典生成工具",
            font=('微软雅黑', 16, 'bold'),
            bootstyle=PRIMARY
        )
        title_label.pack(pady=(0, 20))

        # 作者信息
        author_label = ttk.Label(
            main_frame,
            text="作者：sundi@k1-energy.com",
            font=('微软雅黑', 12)
        )
        author_label.pack(pady=(0, 10))

        # 版本信息
        version_label = ttk.Label(
            main_frame,
            text="版本：1.0.0",
            font=('微软雅黑', 10),
            bootstyle=SECONDARY
        )
        version_label.pack(pady=(0, 20))

        # 关闭按钮
        close_button = ttk.Button(
            main_frame,
            text="确定",
            command=about_window.destroy,
            bootstyle=PRIMARY,
            width=15
        )
        close_button.pack()

        # 设置焦点
        about_window.focus_set()
        about_window.grab_set()  # 模态窗口

    def select_path(self):
        """选择输出路径"""
        path = filedialog.askdirectory(title="选择保存路径", initialdir=self.output_path)
        if path:
            self.output_path = path
            self.path_entry.delete(0, tk.END)
            self.path_entry.insert(0, path)

    def refresh_databases(self):
        """刷新数据库列表"""
        # 验证前置配置
        host = self.host_entry.get().strip()
        user = self.user_entry.get().strip()
        password = self.password_entry.get()
        charset = self.charset_entry.get().strip() or "utf8mb4"

        if not host:
            messagebox.showerror("错误", "请输入主机地址")
            return
        if not user:
            messagebox.showerror("错误", "请输入用户名")
            return

        # 禁用按钮
        self.refresh_db_button.config(state=tk.DISABLED, text="⏳ 连接中...")
        self.update_status("正在连接数据库...", "blue")

        # 在新线程中执行
        def fetch_databases():
            try:
                databases = get_databases(host, user, password, charset)
                if databases:
                    # 在主线程中更新UI
                    self.root.after(0, lambda: self.update_database_list(databases))
                    self.root.after(0, lambda: self.update_status(f"找到 {len(databases)} 个数据库", "green"))
                else:
                    self.root.after(0, lambda: messagebox.showwarning("警告", "未找到可用数据库"))
                    self.root.after(0, lambda: self.update_status("就绪", "green"))
            except pymysql.Error as e:
                self.root.after(0, lambda: messagebox.showerror("错误", f"连接数据库失败：\n{str(e)}"))
                self.root.after(0, lambda: self.update_status("连接失败", "red"))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", f"获取数据库列表失败：\n{str(e)}"))
                self.root.after(0, lambda: self.update_status("获取失败", "red"))
            finally:
                self.root.after(0, lambda: self.refresh_db_button.config(state=tk.NORMAL, text="🔄 刷新列表"))

        thread = threading.Thread(target=fetch_databases, daemon=True)
        thread.start()

    def update_database_list(self, databases):
        """更新数据库下拉列表"""
        self.database_combo['values'] = databases
        if databases:
            self.database_combo.current(0)

    def validate_inputs(self):
        """验证输入"""
        if not self.host_entry.get().strip():
            messagebox.showerror("错误", "请输入主机地址")
            return False
        if not self.user_entry.get().strip():
            messagebox.showerror("错误", "请输入用户名")
            return False
        if not self.database_combo.get().strip():
            messagebox.showerror("错误", "请选择数据库名")
            return False
        if not self.path_entry.get().strip():
            messagebox.showerror("错误", "请选择保存路径")
            return False
        if not os.path.exists(self.path_entry.get().strip()):
            messagebox.showerror("错误", "保存路径不存在")
            return False
        return True

    def update_status(self, message, color="black"):
        """更新状态"""
        # 根据颜色选择样式和前缀
        if color == "green":
            bootstyle = SUCCESS
            prefix = "✓ "
        elif color == "blue":
            bootstyle = INFO
            prefix = "⏳ "
        elif color == "red":
            bootstyle = DANGER
            prefix = "✗ "
        else:
            bootstyle = SUCCESS
            prefix = ""

        # 更新标签样式和文本
        self.status_label.configure(bootstyle=bootstyle, text=prefix + message)
        self.root.update()

    def generate_dict(self):
        """生成数据字典（在后台线程中执行）"""
        try:
            # 获取配置
            host = self.host_entry.get().strip()
            user = self.user_entry.get().strip()
            password = self.password_entry.get()
            database = self.database_combo.get().strip()
            charset = self.charset_entry.get().strip() or "utf8mb4"
            output_path = self.path_entry.get().strip()

            # 连接数据库
            self.update_status("正在连接数据库...", "blue")
            conn = connect_db(host, user, password, database, charset)
            cursor = conn.cursor()

            # 获取表列表
            self.update_status("正在获取表列表...", "blue")
            tables = get_tables(cursor, database)

            if not tables:
                messagebox.showwarning("警告", "数据库中没有找到表")
                cursor.close()
                conn.close()
                self.update_status("就绪", "green")
                self.generate_button.config(state=tk.NORMAL)
                return

            # 获取选择的格式
            format_type = self.format_combo.get()

            # 生成字典
            self.update_status(f"正在生成数据字典（共{len(tables)}个表）...", "blue")
            if format_type == "Markdown格式":
                file_path = generate_markdown(cursor, database, tables, output_path)
            elif format_type == "Word文档格式":
                file_path = generate_doc(cursor, database, tables, output_path)
            else:
                file_path = generate_markdown(cursor, database, tables, output_path)

            # 关闭连接
            cursor.close()
            conn.close()

            # 完成
            self.update_status("生成完成！", "green")
            messagebox.showinfo("成功", f"✨ 数据字典生成完成！\n\n📄 文件保存位置：\n{file_path}")

        except pymysql.Error as e:
            self.update_status("数据库连接失败", "red")
            messagebox.showerror("错误", f"数据库操作失败：\n{str(e)}")
        except Exception as e:
            self.update_status("生成失败", "red")
            messagebox.showerror("错误", f"生成失败：\n{str(e)}")
        finally:
            self.generate_button.config(state=tk.NORMAL, text="✨ 生成数据字典")

    def start_generate(self):
        """开始生成（在新线程中执行）"""
        if not self.validate_inputs():
            return

        self.generate_button.config(state=tk.DISABLED, text="⏳ 正在生成...")
        self.update_status("准备中...", "blue")

        # 在新线程中执行，避免界面卡顿
        thread = threading.Thread(target=self.generate_dict, daemon=True)
        thread.start()


if __name__ == "__main__":
    # 使用 ttkbootstrap 创建窗口，应用现代化主题
    root = ttk.Window(themename="cosmo")  # 可选主题: cosmo, flatly, litera, minty, pulse, sandstone, united, yeti
    app = DatabaseDictGenerator(root)
    root.mainloop()